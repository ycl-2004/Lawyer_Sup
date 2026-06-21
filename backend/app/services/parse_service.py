"""文件解析服务。

轻量原则：txt/md 零依赖必可用；PDF/docx 依赖可选安装，
缺库或解析失败时返回明确错误，由前端引导"手动粘贴文本"兜底——
保证小白用户永远有一条能走通的路。
"""
from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path


@dataclass
class ParseResult:
    text: str
    pages: int
    file_type: str
    ocr_status: str  # not_needed | pending
    error: str | None = None

    @property
    def ok(self) -> bool:
        return self.error is None and bool(self.text.strip())


TEXT_EXTS = {"txt", "md", "markdown", "csv", "log"}
IMAGE_EXTS = {"png", "jpg", "jpeg", "webp", "bmp", "heic"}


def _decode(data: bytes) -> str:
    for enc in ("utf-8", "gb18030", "utf-16"):
        try:
            return data.decode(enc)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def parse_file(filename: str, data: bytes) -> ParseResult:
    ext = Path(filename).suffix.lstrip(".").lower()

    if not data:
        return ParseResult(
            text="", pages=1, file_type=ext or "unknown", ocr_status="not_needed",
            error="上传的文件为空（0 字节），无可解析内容。请确认文件或使用「手动粘贴文本」。",
        )

    if ext in TEXT_EXTS:
        text = _decode(data)
        if not text.strip():
            return ParseResult(
                text="", pages=1, file_type=ext, ocr_status="not_needed",
                error="文件不含可提取文本（疑似空文件或全空白）。请使用「手动粘贴文本」。",
            )
        return ParseResult(text=text, pages=max(1, text.count("\n") // 40 + 1),
                           file_type=ext, ocr_status="not_needed")

    if ext == "pdf":
        try:
            import fitz  # PyMuPDF，可选依赖
        except ImportError:
            return ParseResult(
                text="", pages=1, file_type=ext, ocr_status="not_needed",
                error="服务器未安装 PDF 解析组件（pip install pymupdf）。可点击该材料使用「手动粘贴文本」兜底。",
            )
        try:
            doc = fitz.open(stream=data, filetype="pdf")
            pages = doc.page_count
            text = "\n".join(page.get_text() for page in doc)
            doc.close()
            if not text.strip():
                return ParseResult(
                    text="", pages=pages, file_type=ext, ocr_status="pending",
                    error="PDF 无可提取文本（疑似扫描件）。OCR 尚未接入，请使用「手动粘贴文本」。",
                )
            return ParseResult(text=text, pages=pages, file_type=ext, ocr_status="not_needed")
        except Exception as exc:  # noqa: BLE001
            return ParseResult(text="", pages=1, file_type=ext, ocr_status="not_needed",
                               error=f"PDF 解析失败：{exc}。请使用「手动粘贴文本」。")

    if ext == "docx":
        try:
            from docx import Document  # python-docx，可选依赖
        except ImportError:
            return ParseResult(
                text="", pages=1, file_type=ext, ocr_status="not_needed",
                error="服务器未安装 docx 解析组件（pip install python-docx）。请使用「手动粘贴文本」。",
            )
        try:
            d = Document(BytesIO(data))
            text = "\n".join(p.text for p in d.paragraphs)
            return ParseResult(text=text, pages=max(1, len(d.paragraphs) // 25 + 1),
                               file_type=ext, ocr_status="not_needed")
        except Exception as exc:  # noqa: BLE001
            return ParseResult(text="", pages=1, file_type=ext, ocr_status="not_needed",
                               error=f"docx 解析失败：{exc}。请使用「手动粘贴文本」。")

    if ext in IMAGE_EXTS:
        return ParseResult(
            text="", pages=1, file_type=ext, ocr_status="pending",
            error="图片 OCR 尚未接入（v1 范围外）。请使用「手动粘贴文本」录入图中内容。",
        )

    return ParseResult(
        text="", pages=1, file_type=ext or "unknown", ocr_status="not_needed",
        error=f"暂不支持的文件类型（.{ext}）。支持：PDF、docx、txt/md；其他请粘贴文本。",
    )
